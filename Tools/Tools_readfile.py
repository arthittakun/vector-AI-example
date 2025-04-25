from io import BytesIO
from pypdf import PdfReader
import docx
import pandas as pd
import json
import markdown
import email
from bs4 import BeautifulSoup
import numpy as np
from PIL import Image
import os
import logging
import traceback
from Tools.Tools_media_processor import MediaProcessor

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('file_processing.log'),
        logging.StreamHandler()
    ]
)

# Optional imports with fallback
try:
    import speech_recognition as sr
    SPEECH_RECOGNITION_AVAILABLE = True
except ImportError:
    SPEECH_RECOGNITION_AVAILABLE = False
    print("Warning: speech_recognition not available. Audio processing will be disabled.")

from Model.Model_Vector_DB import VectorDB
from concurrent.futures import ThreadPoolExecutor
import threading
from Tools.Tools_document import Document

progress_lock = threading.Lock()
upload_progress = {}

class Tools_readfile:
    # Maximum file sizes
    MAX_FILE_SIZES = {
        'pdf': 10 * 1024 * 1024,    # 10MB
        'txt': 5 * 1024 * 1024,     # 5MB
        'docx': 8 * 1024 * 1024,    # 8MB
        'doc': 8 * 1024 * 1024,     # 8MB
        'csv': 5 * 1024 * 1024,     # 5MB
        'xlsx': 8 * 1024 * 1024,    # 8MB
        'json': 5 * 1024 * 1024,    # 5MB
        'html': 5 * 1024 * 1024,    # 5MB
        'md': 5 * 1024 * 1024,      # 5MB
        'eml': 5 * 1024 * 1024,     # 5MB
        'jpg': 5 * 1024 * 1024,     # 5MB
        'png': 5 * 1024 * 1024,     # 5MB
        'mp3': 20 * 1024 * 1024,    # 20MB
        'wav': 20 * 1024 * 1024,    # 20MB
        'mp4': 50 * 1024 * 1024,    # 50MB
    }

    def __init__(self):
        self.db = VectorDB()
    
    @staticmethod
    def update_progress(self, file_id, status, message):
        with progress_lock:
            upload_progress[file_id] = {'status': status, 'message': message}

    @staticmethod
    def chunk_text(text, chunk_size=250):
        words = text.split()
        chunks = []
        total_words = len(words)
        
        for i in range(0, total_words, chunk_size):
            chunk = words[i:i + chunk_size]
            chunks.append(' '.join(chunk))
        return chunks

    def process_txt(self, file_content, filename):
        """Process text files"""
        try:
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            text = file_content.decode('latin-1')
        return self._process_text_content(text, filename, 'txt')

    def process_pdf(self, file_content, filename):
        """Process PDF files"""
        pdf_stream = BytesIO(file_content)
        reader = PdfReader(pdf_stream)
        all_text = ""
        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                all_text += text + "\n"
        
        return self._process_text_content(all_text, filename, 'pdf')

    def process_docx(self, file_content, filename):
        """Process Word documents"""
        doc_stream = BytesIO(file_content)
        doc = docx.Document(doc_stream)
        all_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return self._process_text_content(all_text, filename, 'docx')

    def process_excel(self, file_content, filename):
        """Process Excel files"""
        excel_stream = BytesIO(file_content)
        df = pd.read_excel(excel_stream)
        text_content = df.to_string(index=False)
        return self._process_text_content(text_content, filename, 'excel')

    def process_csv(self, file_content, filename):
        """Process CSV files"""
        csv_stream = BytesIO(file_content)
        df = pd.read_csv(csv_stream)
        text_content = df.to_string(index=False)
        return self._process_text_content(text_content, filename, 'csv')

    def process_json(self, file_content, filename):
        """Process JSON files"""
        try:
            json_data = json.loads(file_content.decode('utf-8'))
            text_content = json.dumps(json_data, indent=2, ensure_ascii=False)
            return self._process_text_content(text_content, filename, 'json')
        except UnicodeDecodeError:
            json_stream = BytesIO(file_content)
            text_content = json_stream.read().decode('utf-8')
            return self._process_text_content(text_content, filename, 'json')

    def process_html(self, file_content, filename):
        """Process HTML files"""
        try:
            html_text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            html_text = file_content.decode('latin-1')
        
        soup = BeautifulSoup(html_text, 'html.parser')
        text_content = soup.get_text(separator='\n', strip=True)
        return self._process_text_content(text_content, filename, 'html')

    def process_markdown(self, file_content, filename):
        """Process Markdown files"""
        try:
            md_text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            md_text = file_content.decode('latin-1')
            
        html = markdown.markdown(md_text)
        soup = BeautifulSoup(html, 'html.parser')
        text_content = soup.get_text(separator='\n', strip=True)
        return self._process_text_content(text_content, filename, 'markdown')

    def process_email(self, file_content, filename):
        """Process email files"""
        try:
            msg = email.message_from_bytes(file_content)
            text_parts = []
            
            # Add headers
            text_parts.append(f"Subject: {msg.get('subject', 'No Subject')}")
            text_parts.append(f"From: {msg.get('from', 'Unknown')}")
            text_parts.append(f"To: {msg.get('to', 'Unknown')}")
            
            # Process body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        payload = part.get_payload(decode=True)
                        if payload:
                            try:
                                text_parts.append(payload.decode('utf-8'))
                            except UnicodeDecodeError:
                                text_parts.append(payload.decode('latin-1'))
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    try:
                        text_parts.append(payload.decode('utf-8'))
                    except UnicodeDecodeError:
                        text_parts.append(payload.decode('latin-1'))
            
            text_content = "\n\n".join(text_parts)
            return self._process_text_content(text_content, filename, 'email')
        except Exception as e:
            logging.error(f"Error processing email: {str(e)}")
            return False

    def process_audio(self, file_content, filename):
        """Process audio files (MP3, WAV)"""
        temp_file = None
        try:
            os.makedirs('temp', exist_ok=True)
            temp_file = os.path.join('temp', f"temp_{filename}")
            
            with open(temp_file, 'wb') as f:
                f.write(file_content)

            text_content = MediaProcessor.process_media_file(temp_file)
            return self._process_text_content(text_content, filename, 'audio')

        except Exception as e:
            logging.error(f"Error processing audio file {filename}: {str(e)}")
            return False
        finally:
            if temp_file and os.path.exists(temp_file):
                os.remove(temp_file)

    def process_video(self, file_content, filename):
        """Process video files (MP4)"""
        temp_file = None
        try:
            os.makedirs('temp', exist_ok=True)
            temp_file = os.path.join('temp', f"temp_{filename}")
            
            with open(temp_file, 'wb') as f:
                f.write(file_content)

            text_content = MediaProcessor.process_media_file(temp_file)
            return self._process_text_content(text_content, filename, 'video')

        except Exception as e:
            logging.error(f"Error processing video file {filename}: {str(e)}")
            return False
        finally:
            if temp_file and os.path.exists(temp_file):
                os.remove(temp_file)

    def _process_text_content(self, text_content, filename, file_type):
        """Enhanced text processing with validation"""
        try:
            if not text_content or not text_content.strip():
                raise ValueError("Empty text content after processing")

            chunks = self.chunk_text(text_content)
            if not chunks:
                raise ValueError("No chunks created from content")

            for chunk_id, chunk in enumerate(chunks, 1):
                if not chunk.strip():
                    logging.warning(f"Empty chunk found in {filename} at position {chunk_id}")
                    continue
                    
                doc = Document(
                    page_content=chunk,
                    metadata={
                        'source': filename,
                        'chunk_id': chunk_id,
                        'type': file_type,
                        'total_chunks': len(chunks)
                    }
                )
                self.db.add_document(doc)
                
            logging.info(f"Successfully processed {len(chunks)} chunks from {filename}")
            return True

        except Exception as e:
            logging.error(f"Error in _process_text_content for {filename}: {str(e)}")
            logging.error(f"Traceback: {traceback.format_exc()}")
            return False

    def get_file_processor(self, file_extension: str):
        """Get appropriate file processor based on extension"""
        processors = {
            'txt': self.process_txt,
            'pdf': self.process_pdf,
            'docx': self.process_docx,
            'doc': self.process_docx,
            'csv': self.process_csv,
            'xlsx': self.process_excel,
            'xls': self.process_excel,
            'json': self.process_json,
            'html': self.process_html,
            'md': self.process_markdown,
            'eml': self.process_email,
            'mp3': self.process_audio,
            'wav': self.process_audio,
            'mp4': self.process_video
        }
        return processors.get(file_extension.lower())

    def process_file(self, file_content, filename, file_type):
        """Process file with detailed error handling"""
        processor = self.get_file_processor(file_type)
        if not processor:
            logging.error(f"No processor found for file type: {file_type}")
            return False

        try:
            logging.info(f"Starting to process {filename} of type {file_type}")
            
            # Validate file content
            if not file_content:
                raise ValueError("Empty file content")
                
            # Check file size
            if len(file_content) > self.MAX_FILE_SIZES.get(file_type, 5 * 1024 * 1024):
                raise ValueError(f"File too large for type {file_type}")
                
            # Process the file
            result = processor(file_content, filename)
            
            if result:
                logging.info(f"Successfully processed {filename}")
                return True
            else:
                logging.error(f"Processing failed for {filename}")
                return False

        except Exception as e:
            logging.error(f"Error processing {filename}: {str(e)}")
            logging.error(f"Traceback: {traceback.format_exc()}")
            return False

    def upload_to_vector(self, content, filename, content_type='text'):
        """Enhanced upload method with validation"""
        try:
            logging.info(f"Starting upload: {filename} ({content_type})")
            
            if not content:
                raise ValueError("Empty content provided")
                
            result = self.process_file(content, filename, content_type)
            
            if result:
                logging.info(f"Successfully uploaded {filename} to vector DB")
                return True
            else:
                logging.error(f"Failed to upload {filename}")
                return False

        except Exception as e:
            logging.error(f"Error in upload_to_vector: {str(e)}")
            logging.error(f"Traceback: {traceback.format_exc()}")
            return False


